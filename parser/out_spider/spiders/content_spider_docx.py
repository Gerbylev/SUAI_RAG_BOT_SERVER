import os
import json
from docx_parser import DocumentParser
from docx import Document
import requests
from io import BytesIO


# Основной метод для извлечения текста из документа
def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        text_data = []

        # Извлекаем текст из параграфов
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                text_data.append(para_text)

        # Извлекаем текст из таблиц
        table_data = extract_table_data(file_path)  # Используем второй метод для таблиц
        for table_row in table_data:
            text_data.append(" | ".join(table_row))  # Склеиваем данные ячеек таблицы

        # Объединяем все извлеченные данные в одну строку
        merge_text = " ".join(text_data)
        return merge_text

    except Exception as e:
        print(f"Ошибка при обработке {os.path.basename(file_path)}: {e}")
        return []

# Метод для извлечения данных из таблиц
def extract_table_data(docx_path):
    doc = Document(docx_path)
    table_data = []

    for table in doc.tables:
        for row in table.rows:
            row_data = []

            for idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()

                # Если в ячейке несколько строк, например, во втором столбце
                if '\n' in cell_text:
                    cell_text = " ".join(cell_text.split('\n')).strip()

                row_data.append(cell_text)
            table_data.append(row_data)

    return table_data

# Функция для скачивания документа по URL
def download_docx(url):
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers)
        response.raise_for_status()  # Проверка на ошибки HTTP
        docx_file = BytesIO(response.content)  # Сохраняем контент в байтовый поток
        return docx_file
    except Exception as e:
        print(f"Ошибка при скачивании {url}: {e}")
        return None


# Основной процесс для обработки списка ссылок
def process_docx_links(links):
    result = {}

    for url in links:
        if not url.lower().endswith('.docx'):
            continue

        print(f"Обрабатывается: {url}")

        # Скачиваем документ
        docx_file = download_docx(url)
        if docx_file:
            # Извлекаем текст
            text_content = extract_text_from_docx(docx_file)
            result[url] = text_content

    return result

# Сохранение данных в JSON файл
def save_to_json(data, filename='parsed_docx.json'):
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# Чтение списка ссылок из файла
def read_links_from_file(filename):
    with open(filename, 'r') as file:
        links = file.readlines()
    return [link.strip() for link in links]

if __name__ == "__main__":
    links = read_links_from_file('parser\out_spider\spiders\links.csv')

    parsed_data = process_docx_links(links)

    output_json = 'parsed_docx.json'
    save_to_json(parsed_data, output_json)

    # Статистика
    total_files = len(parsed_data)
    total_text_elements = sum(len(texts) for texts in parsed_data.values())
    print(f"Обработано ссылок: {total_files}")
    print(f"Всего текстовых элементов: {total_text_elements}")